from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from fastapi.concurrency import run_in_threadpool
import json
from typing import List, Dict, Optional
import uvicorn
from docx import Document
import os
from abc import ABC, abstractmethod
from fastapi.middleware.cors import CORSMiddleware
import asyncio
from asyncio import Queue
import uuid


app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

templates = Jinja2Templates(directory="templates")

# Create templates directory if it doesn't exist
os.makedirs("templates", exist_ok=True)


# Store active SSE clients
clients: Dict[str, Queue] = {}


async def add_client() -> str:
    """Add a new SSE client and return its ID."""
    client_id = str(uuid.uuid4())
    clients[client_id] = Queue()
    return client_id


async def remove_client(client_id: str):
    """Remove an SSE client."""
    if client_id in clients:
        del clients[client_id]


async def broadcast_message(message: str):
    """Broadcast a message to all connected clients."""
    for queue in clients.values():
        await queue.put(json.dumps({"type": "log", "message": message}))


async def broadcast_issue(issue: Dict):
    """Broadcast a detected issue to all connected clients."""
    for queue in clients.values():
        await queue.put(json.dumps({"type": "issue", "data": issue}))


@app.get("/sse")
async def sse_endpoint():
    """SSE endpoint for client connections."""
    client_id = await add_client()
    
    async def event_generator():
        try:
            while True:
                if client_id in clients:
                    message = await clients[client_id].get()
                    yield f"data: {message}\n\n"
                else:
                    break
                await asyncio.sleep(0.1)
        except asyncio.CancelledError:
            await remove_client(client_id)
        except Exception as e:
            print(f"SSE Error: {str(e)}")
            await remove_client(client_id)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


class LineBreakDetector(ABC):
    @abstractmethod
    async def detect(self, text1: str, text2: str) -> bool:
        pass


class RuleBasedDetector(LineBreakDetector):
    async def detect(self, text1: str, text2: str) -> bool:
        if not text1 or not text2:
            return False

        # Check if text1 ends with kanji or kana
        last_char = text1[-1]
        is_kanji_or_kana = (
            "\u4e00" <= last_char <= "\u9fff"  # Kanji
            or "\u3040" <= last_char <= "\u309f"  # Hiragana
            or "\u30a0" <= last_char <= "\u30ff"
        )  # Katakana

        # Check if text1 doesn't end with common sentence endings
        no_sentence_end = not any(
            text1.endswith(end) for end in ["。", "！", "？", ".", "!", "?"]
        )

        # Check if text1 is too short to be a sentence (just caption)
        if len(text1) < 20:
            return False

        # Check if text starts with common section markers
        for text in [text1, text2]:
            # Check for Chinese/Japanese section markers
            if text[0] in ["第", "序", "附"]:
                # Check for numbers after markers (both Chinese and Arabic)
                if len(text) > 1 and (
                    text[1]
                    in [
                        "一",
                        "二",
                        "三",
                        "四",
                        "五",
                        "六",
                        "七",
                        "八",
                        "九",
                        "十",
                        "零",
                        "壱",
                        "弐",
                        "参",
                        "肆",
                        "伍",
                        "陸",
                        "柒",
                        "捌",
                        "玖",
                        "拾",
                    ]
                    or text[1].isdigit()
                ):
                    return False

            # Check for English section markers
            for prefix in [
                "Chapter",
                "chapter",
                "Section",
                "section",
                "Part",
                "part",
                "Article",
                "article",
                "Clause",
                "clause",
                "目",
                "Index",
                "index",
                "Table",
                "table",
            ]:
                if text.strip().startswith(prefix):
                    return False

            # Check for numbers at start (handling all types of spaces)
            first_non_space = ""
            for char in text:
                if not is_space_character(char):
                    first_non_space = char
                    break

            # Check if first non-space character is a number (any format)
            if first_non_space:
                # Check Arabic numerals
                if first_non_space.isdigit():
                    return False

                # Check Chinese/Japanese numerals
                chinese_numbers = [
                    "零",
                    "一",
                    "二",
                    "三",
                    "四",
                    "五",
                    "六",
                    "七",
                    "八",
                    "九",
                    "十",
                    "壱",
                    "弐",
                    "参",
                    "肆",
                    "伍",
                    "陸",
                    "柒",
                    "捌",
                    "玖",
                    "拾",
                ]
                if first_non_space in chinese_numbers:
                    return False

        # Broadcast detection result
        await broadcast_message(
            f"Rule-based detector: Checking connection between '{text1}' and '{text2}'"
        )
        result = is_kanji_or_kana and no_sentence_end
        await broadcast_message(f"Rule-based detector result: {result}")

        return result


def is_space_character(char):
    """Check if a character is any type of space"""
    return char in [
        " ",  # Regular space
        "\t",  # Tab
        "\n",  # Newline
        "\r",  # Carriage return
        "\f",  # Form feed
        "\v",  # Vertical tab
        "\u00A0",  # Non-breaking space
        "\u2000",  # En quad
        "\u2001",  # Em quad
        "\u2002",  # En space
        "\u2003",  # Em space
        "\u2004",  # Three-per-em space
        "\u2005",  # Four-per-em space
        "\u2006",  # Six-per-em space
        "\u2007",  # Figure space
        "\u2008",  # Punctuation space
        "\u2009",  # Thin space
        "\u200A",  # Hair space
        "\u200B",  # Zero-width space
        "\u3000",  # Ideographic space (CJK)
    ]


class LLMDetector(LineBreakDetector):
    def __init__(self, llm_config: Dict):
        from litellm import completion

        # Set environment variables from config
        if "envVars" in llm_config:
            for env_var in llm_config.get("envVars", []):
                if env_var["name"].strip() == "":
                    continue
                os.environ[env_var["name"]] = env_var["value"]

        self.llm = completion
        self.model = llm_config.get("modelName", "openai/gpt-3.5-turbo")
        
        # Number of characters to take from end/beginning of paragraphs
        self.context_length = 50
        
        self.prompt_template = (
            "Are these two text segments **contiguous** part of the same sentence? Answer only 'yes' or 'no'. "
            "Segment 1: {text1} Segment 2: {text2}"
        )

    def _get_relevant_segments(self, text1: str, text2: str) -> tuple[str, str]:
        """Extract relevant portions from the paragraphs."""
        # Get end of first paragraph
        text1 = text1.strip()
        if len(text1) > self.context_length:
            text1 = "..." + text1[-self.context_length:]
            
        # Get start of second paragraph
        text2 = text2.strip()
        if len(text2) > self.context_length:
            text2 = text2[:self.context_length] + "..."
            
        return text1, text2

    async def detect(self, text1: str, text2: str) -> bool:
        from loguru import logger

        # Early return for empty strings
        if text1.strip() == "" or text2.strip() == "":
            return False

        # Get relevant segments of text
        segment1, segment2 = self._get_relevant_segments(text1, text2)
        
        # First broadcast
        await broadcast_message(
            f"LLM detecting connection between segments: '{segment1}' and '{segment2}'"
        )
        prompt = self.prompt_template.format(text1=segment1, text2=segment2)

        try:
            # Try LLM call
            response = await run_in_threadpool(
                lambda: self.llm(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                    max_tokens=10,
                )
            )
            logger.debug("LLM call completed")

            result = response.choices[0].message.content.lower().strip()
            await broadcast_message(f"LLM response: {result}")
            return result == "yes"

        except Exception as e:
            error_msg = f"LLM API Error: {str(e)}"
            logger.critical(error_msg)
            await broadcast_message(error_msg)
            return False


class DocumentProcessor:
    def __init__(
        self, detector_config: Optional[Dict] = None, llm_config: Optional[Dict] = None
    ):
        self.detectors = []

        if detector_config:
            if detector_config.get("ruleBased", False):
                self.detectors.append(RuleBasedDetector())
            if detector_config.get("llm", True):
                self.detectors.append(LLMDetector(llm_config or {}))

    async def process_document(self, file_path: str) -> List[Dict]:
        await broadcast_message(f"Processing document: {file_path}")
        doc = Document(file_path)
        issues = []

        for para_idx, para in enumerate(doc.paragraphs[:-1]):
            next_para = doc.paragraphs[para_idx + 1]

            for detector in self.detectors:
                if await detector.detect(para.text, next_para.text):
                    issue = {
                        "paragraph1": para.text,
                        "paragraph2": next_para.text,
                        "index": para_idx,
                        "detector": detector.__class__.__name__,
                    }
                    issues.append(issue)
                    await broadcast_message(f"Found issue at paragraph {para_idx}")
                    await broadcast_issue(issue)

        return issues

    async def fix_document(self, file_path: str, issues: List[Dict]) -> str:
        await broadcast_message("Starting document fix...")
        doc = Document(file_path)
        fixed_doc = Document()  # Create a new document
        fixed_doc_path = f"fixed_{os.path.basename(file_path)}"

        # Sort issues by index to process them in order
        sorted_issues = sorted(issues, key=lambda x: x["index"])
        issue_indices = {issue["index"] for issue in sorted_issues}

        # Copy styles from original document
        for style in doc.styles:
            if style.name not in fixed_doc.styles:
                fixed_doc.styles.add_style(style.name, style.type, style.base_style)

        # Process paragraphs
        i = 0
        while i < len(doc.paragraphs):
            if i in issue_indices:
                # Merge this paragraph with the next one
                current_para = doc.paragraphs[i]
                next_para = doc.paragraphs[i + 1]

                # Create new paragraph with style of first paragraph
                new_para = fixed_doc.add_paragraph()
                if current_para.style:
                    new_para.style = current_para.style

                # Copy runs from first paragraph
                for run in current_para.runs:
                    new_run = new_para.add_run(run.text)
                    # Copy run formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color:
                        new_run.font.color.rgb = run.font.color.rgb

                # Copy runs from second paragraph
                for run in next_para.runs:
                    new_run = new_para.add_run(run.text)
                    # Copy run formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color:
                        new_run.font.color.rgb = run.font.color.rgb

                await broadcast_message(f"Merged paragraphs at index {i}")
                i += 2  # Skip next paragraph since we merged it
            else:
                # Copy paragraph as is
                current_para = doc.paragraphs[i]
                new_para = fixed_doc.add_paragraph()

                # Copy paragraph style
                if current_para.style:
                    new_para.style = current_para.style

                # Copy runs with formatting
                for run in current_para.runs:
                    new_run = new_para.add_run(run.text)
                    # Copy run formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color:
                        new_run.font.color.rgb = run.font.color.rgb

                i += 1

        fixed_doc.save(fixed_doc_path)
        await broadcast_message(f"Document fixed and saved as: {fixed_doc_path}")
        return fixed_doc_path


@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/preview")
async def preview_document(file: UploadFile = File(...)):
    temp_path = f"temp_{file.filename}"
    try:
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)

        doc = Document(temp_path)
        paragraphs = [para.text for para in doc.paragraphs]
        await broadcast_message(f"Generated preview for {file.filename}")

        return {"paragraphs": paragraphs}
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


@app.post("/api/process")
async def process_document(
    file: UploadFile = File(...),
    detectors: str = Form(...),
    llmConfig: str = Form(...),
    issues: Optional[str] = Form(None),
):
    detector_config = json.loads(detectors)
    llm_config = json.loads(llmConfig)
    temp_path = f"temp_{file.filename}"

    try:
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)

        processor = DocumentProcessor(detector_config, llm_config)

        if issues:
            fixed_path = await processor.fix_document(temp_path, json.loads(issues))
            return {"filename": fixed_path}

        issues = await processor.process_document(temp_path)
        return {"issues": issues}

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    return FileResponse(
        path=filename,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=28000)
